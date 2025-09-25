from docx import Document as DocxDocument
from docx.shared import Inches
from datetime import datetime
from pathlib import Path
from typing import Dict, List

def generate_warning_document(employee_data: Dict[str, str], output_dir: Path, template_type: str = "B1") -> None:
    """生成基于 Annex B1, B2, E, F, G, H, C 或 D 模板的文档"""
    if not isinstance(template_type, str):
        raise ValueError("template_type must be a string")
    template_type = template_type.upper()

    doc = DocxDocument()

    if template_type in ["B1", "B2"]:
        # 设置警告模板参数
        is_first_warning = template_type == "B1"
        annexure = "B1" if is_first_warning else "B2"
        warning_title = "WRITTEN WARNING" if is_first_warning else "FINAL WRITTEN WARNING"
        warning_period = "6 months" if is_first_warning else "12 months"
        
        # 添加标题
        doc.add_heading(f"DISCIPLINARY CODES AND PROCEDURES ANNEXURE '{annexure}'", 0)
        doc.add_heading(warning_title, 1)
        
        # 员工信息表格
        doc.add_heading("Employee Information", 2)
        table = doc.add_table(rows=3, cols=4)
        table.style = 'Table Grid'
        
        # 第一行
        table.cell(0, 0).text = "Employee Name:"
        table.cell(0, 1).text = employee_data.get("name", "")
        table.cell(0, 2).text = "Date:"
        table.cell(0, 3).text = employee_data.get("date", datetime.now().strftime("%d %B %Y"))
        
        # 第二行
        table.cell(1, 0).text = "Employee ID:"
        table.cell(1, 1).text = employee_data.get("employee_id", "")
        table.cell(1, 2).text = "Job Title:"
        table.cell(1, 3).text = employee_data.get("job_title", "")
        
        # 第三行
        table.cell(2, 0).text = "Manager:"
        table.cell(2, 1).text = employee_data.get("manager", "Darren Britz")
        table.cell(2, 2).text = "Department:"
        table.cell(2, 3).text = employee_data.get("department", "Legal")
        
        # 警告类型
        doc.add_heading("Type of Warning", 2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        warning_type = "First Warning" if is_first_warning else "Final Warning"
        table.cell(0, 0).text = warning_type
        table.cell(0, 1).text = "☑"
        
        # 违法行为类型
        doc.add_heading("Type of Offenses", 2)
        table = doc.add_table(rows=2, cols=6)
        table.style = 'Table Grid'
        
        offense_types = employee_data.get("offense_types", [])
        offense_options = [
            "Tardiness/Leaving Early",
            "Absenteeism",
            "Violation of Policies / Procedures",
            "Substandard Work",
            "Violation of Safety Rules",
            "Rudeness to Customers/Co-workers"
        ]
        
        # 第一行
        table.cell(0, 0).text = "☑" if "Tardiness/Leaving Early" in offense_types else "☐"
        table.cell(0, 1).text = "Tardiness/Leaving Early"
        table.cell(0, 2).text = "☑" if "Absenteeism" in offense_types else "☐"
        table.cell(0, 3).text = "Absenteeism"
        table.cell(0, 4).text = "☑" if "Violation of Policies / Procedures" in offense_types else "☐"
        table.cell(0, 5).text = "Violation of Policies / Procedures"
        
        # 第二行
        table.cell(1, 0).text = "☑" if "Substandard Work" in offense_types else "☐"
        table.cell(1, 1).text = "Substandard Work"
        table.cell(1, 2).text = "☑" if "Violation of Safety Rules" in offense_types else "☐"
        table.cell(1, 3).text = "Violation of Safety Rules"
        table.cell(1, 4).text = "☑" if "Rudeness to Customers/Co-workers" in offense_types else "☐"
        table.cell(1, 5).text = "Rudeness to Customers/Co-workers"
        
        # 其他违法行为
        if "Other" in offense_types:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.cell(0, 0).text = "Other:"
            table.cell(0, 1).text = employee_data.get("other_offense", "")
        
        # 详细信息
        doc.add_heading("Details", 2)
        warning_text = "first written warning" if is_first_warning else "final written warning"
        doc.add_paragraph(
            f"This is a {warning_text} to the employee about his/her failure to perform task(s) or duty without the exercise of due care and attention."
        )
        doc.add_paragraph(f"Description of Infraction: {employee_data.get('infraction', '未提供违规详情')}")
        doc.add_paragraph("Future transgressions by the employee could result in further disciplinary measures.")
        
        # 监督措施
        doc.add_heading("Supervisory Measures", 3)
        doc.add_paragraph(
            f"The employee’s work performance will be carefully supervised and tracked for the duration of this warning, which is {warning_period} from the date herein."
        )
        doc.add_paragraph(
            "1. Performance assessment updates will be required monthly by the immediate line manager and the employee. "
            "The line manager may choose to delegate this role to a co-worker with more direct contact with the employee. "
            "These updates will be required for a three-month period from date of issuance. "
            "These meetings will serve to track progress and manage quality control."
        )
        
        # B2 模板特有的 6 个月后报告要求
        if not is_first_warning:
            doc.add_paragraph(
                "2. After six months, the employee will be required to indicate performance and proof of procedural "
                "considerations monthly in the form of a summary report. This report should indicate computations done, reviewed, and control measures throughout."
            )
        
        doc.add_paragraph(
            "The employee has the opportunity to put in writing reasons for his/her actions/inaction and send this to relevant parties; "
            "however, this will not cancel the stipulations of this issued warning."
        )
        
        # 进一步违规的后果
        doc.add_heading("Consequences of Further Infractions:", 3)
        doc.add_paragraph(
            f"1. This warning is valid for a period of {warning_period} from the date herein."
        )
        doc.add_paragraph(
            "2. The employee may appeal in writing to a higher level of management should he/she not accept the sanction implemented "
            "(an appeal will not stay or suspend the implementation of the sanction) or otherwise confirm acknowledgement and understanding of this written warning."
        )
        
        # 确认收到警告
        doc.add_heading("Acknowledgment of Receipt of Warning", 2)
        doc.add_paragraph(
            "By signing this form, you confirm that you understand the information in this warning. "
            "You also confirm that you and your manager have discussed the warning and a plan for improvement. "
            "Signing this form does not necessarily indicate that you agree with this warning. (You have the right to appeal as indicated above.)"
        )
        doc.add_paragraph(
            "In the event of an appeal the employee must indicate whether the appeal is:\n"
            "1. Against the finding\n"
            "2. Against the sanction\n"
            "3. Against both"
        )
        doc.add_paragraph(
            "Any appeal must be in writing and well-motivated. Such appeal must be made within 3 working days from the date of the implementation of the sanction. "
            "It should not be construed that an appeal in writing guarantees an appeal hearing and an appeal may be dismissed on application should it be found to wanting. "
            "A written appeal not made within 3 working days will be interpreted as that the employee accepts the written warning."
        )
        
        # 签名表格
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "Employee Signature"
        table.cell(0, 1).text = "Date"
        table.cell(1, 0).text = "Manager Signature"
        table.cell(1, 1).text = "Date"
        table.cell(2, 0).text = "Witness Signature (if employee understands warning but refuses to sign)"
        table.cell(2, 1).text = "Date"
        
        # 保存文档
        output_path = output_dir / f"{employee_data['name']}_{employee_data['infraction'].replace(' ', '_')}.docx"
        doc.save(output_path)
        print(f"文档已生成: {output_path}")

    elif template_type == "E":
        # 上诉信模板
        doc.add_heading("DISCIPLINARY CODES AND PROCEDURES ANNEXURE 'E'", 0)
        doc.add_heading("LODGING OF APPEAL", 1)
        
        # 致公司
        doc.add_paragraph(f"TO COMPANY: {employee_data.get('company', '')}")
        
        # 上诉声明
        doc.add_paragraph(
            "In terms of the disciplinary procedures, I wish to lodge an appeal against the decision of the disciplinary enquiry."
        )
        
        # 上诉人姓名
        doc.add_paragraph(f"NAME OF APPELLANT: {employee_data.get('appellant_name', '')}")
        
        # 上诉理由
        doc.add_paragraph(
            "The appeal is lodged on the following grounds (the appropriate areas to be marked with an X)"
        )
        appeal_reasons = employee_data.get("appeal_reasons", [])
        appeal_options = [
            "The penalty imposed was not commensurate with the breach of the disciplinary code.",
            "Disciplinary procedures were not followed.",
            "New or further evidence or witnesses are available, which could bring new facts to bear and which could materially affect the result of the previous hearing."
        ]
        
        for i, reason in enumerate(appeal_options, 1):
            marker = "X" if reason in appeal_reasons else " "
            doc.add_paragraph(f"{i}. {reason} [{marker}]")
        
        # 支持上诉的理由
        doc.add_paragraph("The following is submitted in support of this appeal:")
        doc.add_paragraph(employee_data.get("appeal_details", ""))
        
        # 签名和日期
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "DATE APPEAL LODGED"
        table.cell(0, 1).text = "APPELLANT’S SIGNATURE"
        
        # 保存文档
        output_path = output_dir / f"{employee_data['appellant_name']}_Appeal_Letter.docx"
        doc.save(output_path)
        print(f"文档已生成: {output_path}")

    elif template_type == "F":
        # 辞职信模板
        doc.add_heading("DISCIPLINARY CODES AND PROCEDURES ANNEXURE 'F'", 0)
        doc.add_heading("RESIGNATION BY EMPLOYEE", 1)
        
        # 致公司
        doc.add_paragraph(f"TO COMPANY: {employee_data.get('company', '')}")
        
        # 员工姓名
        doc.add_paragraph(f"NAME OF EMPLOYEE: {employee_data.get('employee_name', '')}")
        
        # 辞去职务
        doc.add_paragraph(f"I hereby resign from my position as {employee_data.get('position', '')}.")
        
        # 最后上班日期
        doc.add_paragraph(f"My last working day shall be {employee_data.get('last_working_day', '')}.")
        
        # 自由意愿声明
        doc.add_paragraph(
            "I certify that I am resigning of my own free will and have not been forced to resign."
        )
        
        # 签名地点和日期
        signed_at = employee_data.get('signed_at', '')
        signed_date = employee_data.get('signed_date', datetime.now().strftime("%d %B %Y"))
        day, month, year = signed_date.split(" ")
        doc.add_paragraph(f"SIGNED AT {signed_at} ON THIS {day} DAY OF {month.upper()} {year}.")
        
        # 辞职原因
        doc.add_paragraph("Reason for resignation:")
        doc.add_paragraph(employee_data.get("resignation_reason", ""))
        
        # 签名部分
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "EMPLOYEE"
        table.cell(0, 1).text = "DATE"
        table.cell(1, 0).text = "EMPLOYER"
        table.cell(1, 1).text = "Resignation accepted:"
        table.cell(2, 0).text = "WITNESS"
        table.cell(2, 1).text = "DATE"
        
        # 备注
        doc.add_paragraph(
            "NOTE: Employee to provide reason for resignation in his/her own handwriting if possible. "
            "Should the Employee be illiterate, ensure that a witness is present when the form is completed "
            "and that the form is explained to the Employee in the presence of said witness."
        )
        
        # 保存文档
        output_path = output_dir / f"{employee_data['employee_name']}_Resignation_Letter.docx"
        doc.save(output_path)
        print(f"文档已生成: {output_path}")

    elif template_type == "G":
        # 最终结算表模板
        doc.add_heading("DISCIPLINARY CODES AND PROCEDURES ANNEXURE 'G'", 0)
        doc.add_heading("FINAL SETTLEMENT", 1)
        
        # 声明
        doc.add_paragraph(
            f"I the undersigned, {employee_data.get('employee_name', '')}, acknowledge receipt of the following in full and final "
            "settlement of all and any claims and/or disputes which may arise now or in the future as a result of the termination "
            "of my contract of employment."
        )
        
        # 金额表格
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "WAGES"
        table.cell(0, 1).text = f"R {employee_data.get('wages', '0.00')}"
        table.cell(1, 0).text = "NOTICE"
        table.cell(1, 1).text = f"R {employee_data.get('notice_pay', '0.00')}"
        table.cell(2, 0).text = "LEAVE PAY"
        table.cell(2, 1).text = f"R {employee_data.get('leave_pay', '0.00')}"
        table.cell(3, 0).text = "SEVERANCE PAY"
        table.cell(3, 1).text = f"R {employee_data.get('severance_pay', '0.00')}"
        table.cell(4, 0).text = "LESS DEDUCTIONS"
        table.cell(4, 1).text = f"R {employee_data.get('deductions', '0.00')}"
        table.cell(5, 0).text = "TOTAL"
        table.cell(5, 1).text = f"R {employee_data.get('total_amount', '0.00')}"
        
        # 签名地点和日期
        signed_at = employee_data.get('signed_at', '')
        signed_date = employee_data.get('signed_date', datetime.now().strftime("%d %B %Y"))
        day, month, year = signed_date.split(" ")
        doc.add_paragraph(f"SIGNED AT {signed_at} ON THIS {day} DAY OF {month.upper()} {year}.")
        
        # 签名表格
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "EMPLOYEE"
        table.cell(0, 1).text = "DATE"
        table.cell(1, 0).text = "WITNESS"
        table.cell(1, 1).text = "DATE"
        table.cell(2, 0).text = "EMPLOYER"
        table.cell(2, 1).text = "DATE"
        
        # 保存文档
        output_path = output_dir / f"{employee_data['employee_name']}_Final_Settlement.docx"
        doc.save(output_path)
        print(f"文档已生成: {output_path}")

    elif template_type == "H":
        # 劳动关系终止表模板
        doc.add_heading("DISCIPLINARY CODES AND PROCEDURES - ANNEXURE 'H'", 0)
        doc.add_heading("TERMINATION OF EMPLOYMENT", 1)
        
        # 收件人（员工）
        doc.add_paragraph(f"ADDRESSEE (Employee): {employee_data.get('employee_name', '')}")
        
        # 终止通知
        termination_date = employee_data.get('termination_date', datetime.now().strftime("%d %B %Y"))
        day, month, year = termination_date.split(" ")
        doc.add_paragraph(
            f"We regret to advise that your employment shall terminate on the {day} day of {month.upper()} {year}. "
            "This action has been necessitated for the following reason/s:"
        )
        doc.add_paragraph(employee_data.get("termination_reason", "No reason provided."))
        
        # 归还公司财产
        doc.add_paragraph("Please arrange for the return of all property of the Employer in your possession.")
        
        # 生效日期
        doc.add_paragraph(f"Date effective: {termination_date}")
        
        # 金额表格
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "Notice period"
        table.cell(0, 1).text = f"R {employee_data.get('notice_period', '0.00')}"
        table.cell(1, 0).text = "Wages due"
        table.cell(1, 1).text = f"R {employee_data.get('wages_due', '0.00')}"
        table.cell(2, 0).text = "Leave pay due"
        table.cell(2, 1).text = f"R {employee_data.get('leave_pay_due', '0.00')}"
        table.cell(3, 0).text = "Add payment in lieu of notice"
        table.cell(3, 1).text = f"R {employee_data.get('payment_in_lieu', '0.00')}"
        table.cell(4, 0).text = "Severance pay if applicable"
        table.cell(4, 1).text = f"R {employee_data.get('severance_pay', '0.00')}"
        table.cell(5, 0).text = "Advance payments outstanding"
        table.cell(5, 1).text = f"R {employee_data.get('advance_payments', '0.00')}"
        table.cell(6, 0).text = "Total amount due"
        table.cell(6, 1).text = f"R {employee_data.get('total_amount_due', '0.00')}"
        
        # 签名
        doc.add_paragraph("Yours faithfully")
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "EMPLOYER"
        table.cell(0, 1).text = "EMPLOYEE"
        
        # 保存文档
        output_path = output_dir / f"{employee_data['employee_name']}_Termination_Letter.docx"
        doc.save(output_path)
        print(f"文档已生成: {output_path}")

    elif template_type == "C":
        # 通知听证会模板
        doc.add_heading("DISCIPLINARY CODES AND PROCEDURES - ANNEXURE 'C'", 0)
        doc.add_heading("NOTIFICATION OF DISCIPLINARY HEARING", 1)
        
        # 收件人和日期
        doc.add_paragraph(f"To alleged offender: {employee_data.get('employee_name', '')}")
        doc.add_paragraph(f"Date: {employee_data.get('notification_date', datetime.now().strftime('%d %B %Y'))}")
        
        # 通知声明
        doc.add_paragraph(
            "You are hereby notified that a formal Disciplinary Hearing will be conducted in respect of the following alleged charges against you: (summary detail)"
        )
        
        # 指控详情
        doc.add_heading("CHARGE:", 2)
        doc.add_paragraph(f"a. On or about {employee_data.get('charge_date', '')}")
        doc.add_paragraph(f"b. {employee_data.get('charge_description', '')}")
        
        # 听证会详情
        doc.add_paragraph("This enquiry will be held:")
        doc.add_paragraph(f"Venue: {employee_data.get('venue', '')}")
        doc.add_paragraph(f"Date: {employee_data.get('hearing_date', '')}")
        doc.add_paragraph(f"Time: {employee_data.get('hearing_time', '')}")
        
        # 员工权利
        doc.add_paragraph("You do have the following rights:")
        rights = [
            "To present your case.",
            "Cross-examination of management and management's witnesses.",
            "The right to be represented at the enquiry by a fellow employee/representative.",
            "The right to call witnesses. (You are to arrange your own witnesses – in cases of Company employees you are to timeously advise management at least 3 working days in advance of the names).",
            "The right to an interpreter. (Timeous notice to Management is required - at least 3 working days in advance)"
        ]
        for i, right in enumerate(rights, 1):
            doc.add_paragraph(f"{chr(96+i)}. {right}")
        
        # 确认声明
        doc.add_paragraph(
            "I fully understand the content of this notification and do not sign this document as an admission of guilt."
        )
        
        # 签名部分
        doc.add_paragraph(f"Signature: {'_'*20}")
        doc.add_paragraph(f"Date: {employee_data.get('signature_date', datetime.now().strftime('%d %B %Y'))}")
        doc.add_paragraph(f"Time: {employee_data.get('signature_time', '')}")
        
        # 备注
        doc.add_paragraph(
            "Note: Should you not attend the Disciplinary Hearing, Management reserves the right to nevertheless continue in your absence "
            "and make a finding and you are therefore URGED to attend the Hearing."
        )
        
        # 保存文档
        output_path = output_dir / f"{employee_data['employee_name']}_Disciplinary_Hearing_Notice.docx"
        doc.save(output_path)
        print(f"文档已生成: {output_path}")

    elif template_type == "D":
        # 纪律报告模板
        doc.add_heading("DISCIPLINARY CODES AND PROCEDURES - ANNEXURE 'D'", 0)
        doc.add_heading("DISCIPLINARY REPORT FORM", 1)
        
        # 地点和日期
        doc.add_paragraph(f"VENUE: {employee_data.get('venue', '')}")
        doc.add_paragraph(f"DATE: {employee_data.get('hearing_date', '')}")
        
        # PRESENT 表格
        doc.add_heading("PRESENT", 2)
        table = doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "CAPACITY"
        table.cell(0, 1).text = "NAME"
        table.cell(0, 2).text = "DESIGNATION"
        table.cell(0, 3).text = "SECTION"
        table.cell(1, 0).text = "Chairperson"
        table.cell(1, 1).text = employee_data.get('chairperson_name', '')
        table.cell(1, 2).text = employee_data.get('chairperson_designation', '')
        table.cell(1, 3).text = employee_data.get('chairperson_section', '')
        table.cell(2, 0).text = "Panel member"
        table.cell(2, 1).text = employee_data.get('panel_member_1_name', '')
        table.cell(2, 2).text = employee_data.get('panel_member_1_designation', '')
        table.cell(2, 3).text = employee_data.get('panel_member_1_section', '')
        table.cell(3, 0).text = "Panel member"
        table.cell(3, 1).text = employee_data.get('panel_member_2_name', '')
        table.cell(3, 2).text = employee_data.get('panel_member_2_designation', '')
        table.cell(3, 3).text = employee_data.get('panel_member_2_section', '')
        
        # 投诉人和员工证人
        complainant_witnesses = employee_data.get('complainant_witnesses', [])
        employee_witnesses = employee_data.get('employee_witnesses', [])
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "COMPLAINANT (Witnesses for Complainant):"
        table.cell(0, 1).text = "EMPLOYEE (Witnesses for Employee):"
        for i in range(3):
            table.cell(i+1, 0).text = complainant_witnesses[i] if i < len(complainant_witnesses) else ""
            table.cell(i+1, 1).text = employee_witnesses[i] if i < len(employee_witnesses) else ""
        
        # 口译员和代表
        doc.add_paragraph(f"Interpreter: {employee_data.get('interpreter', '')}")
        doc.add_paragraph(f"Representative: {employee_data.get('representative', '')}")
        
        # 主席职责
        doc.add_heading("Functions of chairperson:", 2)
        functions = [
            "Introduce those present and state their functions.",
            "Ensure that witnesses are present only while giving their evidence."
        ]
        for i, func in enumerate(functions, 1):
            doc.add_paragraph(f"{i}. {func}")
        
        # 员工权利
        doc.add_heading("Employee's rights (To be read by chairperson)", 2)
        rights = [
            "To present your case.",
            "Cross-examination of management and management's witnesses.",
            "The right to be represented at the enquiry by a fellow employee/representative.",
            "The right to call witnesses. (You are to arrange your own witnesses – in cases of Company employees you are to timeously advise management at least 3 working days in advance of the names).",
            "The right to an interpreter. (Timeous notice to Management is required - at least 3 working days in advance)"
        ]
        for i, right in enumerate(rights, 1):
            doc.add_paragraph(f"{i}. {right}")
        
        # 指控详情
        doc.add_heading("3. Nature of alleged breach or misconduct (Charge, date, time, place and brief description of the incident/s):", 2)
        doc.add_paragraph(employee_data.get("misconduct_details", ""))
        
        # 认罪情况
        doc.add_paragraph(f"4. Plea: The Employee {employee_data.get('plea', 'admits')} the charge/s")
        
        # 调查程序
        doc.add_heading("5. Procedure of enquiry (to be explained by the chairperson)", 2)
        doc.add_paragraph(
            "The evidence of the complainant and his / her witnesses shall be heard first. The Employee and panel may ask questions regarding the evidence. "
            "The Employee and his / her witnesses may then give evidence and they may be questioned by the complainant and panel."
        )
        
        # 已确认事实
        doc.add_heading("FACTS ESTABLISHED", 2)
        doc.add_paragraph(employee_data.get("facts_established", ""))
        
        # 听证会结果
        doc.add_heading("FINDINGS / RESULT OF HEARING: (RECONVENE HEARING)", 2)
        doc.add_paragraph(f"GUILTY / NOT GUILTY: {employee_data.get('hearing_result', 'NOT GUILTY')}")
        
        # 主席行动
        doc.add_heading("Action by chairperson", 2)
        actions = [
            "When evidence has been heard, close the enquiry and advise the complainant and the accused of the date and time to return for a decision. If possible, a decision should be arrived at as soon as possible in order to minimise delays.",
            "Dismiss all witnesses.",
            "Weigh the evidence and come to a decision.",
            "When the enquiry has been reconvened with the accused and complainant being present, the decision must be explained to all parties."
        ]
        for i, action in enumerate(actions, 1):
            doc.add_paragraph(f"{i}. {action}")
        
        # 记录审查
        doc.add_heading("REVIEW OF RECORD", 2)
        doc.add_paragraph(f"Date engaged: {employee_data.get('date_engaged', '')}")
        doc.add_paragraph(f"Years: {employee_data.get('years_service', '')} Months: {employee_data.get('months_service', '')}")
        
        # 先前纪律行动
        doc.add_heading("Previous disciplinary actions", 2)
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "Date"
        table.cell(0, 1).text = "Type of misconduct"
        table.cell(0, 2).text = "Action taken"
        previous_actions = employee_data.get('previous_actions', [])
        for i, action in enumerate(previous_actions[:3], 1):
            table.cell(i, 0).text = action.get('date', '')
            table.cell(i, 1).text = action.get('misconduct_type', '')
            table.cell(i, 2).text = action.get('action_taken', '')
        
        # 减轻因素和处罚建议
        doc.add_heading("MITIGATING FACTORS & REPRESENTATIONS REGARDING PENALTY", 2)
        doc.add_paragraph(employee_data.get("mitigating_factors", ""))
        
        # 投诉人提交
        doc.add_heading("SUBMISSION BY COMPLAINANT", 2)
        doc.add_paragraph(employee_data.get("complainant_submission", ""))
        
        # 主席进一步行动
        doc.add_heading("Action by chairperson", 2)
        further_actions = [
            "Review of record: If the Employee has been found guilty, the Employee's previous valid disciplinary record must now be produced. This must be done in the presence of the Employee and complainant.",
            "The accused's recommended penalty: The accused must be given an opportunity to recommend an appropriate penalty.",
            "After the submission of mitigating factors, the meeting will adjourn to consider."
        ]
        for i, action in enumerate(further_actions, 1):
            doc.add_paragraph(f"{i}. {action}")
        
        # 延期信息
        doc.add_heading("ENQUIRY POSTPONED TO:", 2)
        doc.add_paragraph(f"Date: {employee_data.get('postponed_date', '')}")
        doc.add_paragraph(f"Time: {employee_data.get('postponed_time', '')}")
        doc.add_paragraph(f"Place: {employee_data.get('postponed_place', '')}")
        
        # 重新召开听证会
        doc.add_heading("DISCIPLINARY ENQUIRY RECONVENED", 2)
        doc.add_paragraph(f"Date: {employee_data.get('reconvened_date', '')}")
        doc.add_paragraph(f"Time: {employee_data.get('reconvened_time', '')}")
        doc.add_paragraph(f"Place: {employee_data.get('reconvened_place', '')}")
        
        # 出勤情况
        doc.add_paragraph(f"Employee: {'Present' if employee_data.get('employee_present', 'Present') == 'Present' else 'Absent'}")
        doc.add_paragraph(f"Reason for absence: {employee_data.get('employee_absence_reason', '')}")
        doc.add_paragraph(f"Complainant: {'Present' if employee_data.get('complainant_present', 'Present') == 'Present' else 'Absent'}")
        doc.add_paragraph(f"Reason for absence: {employee_data.get('complainant_absence_reason', '')}")
        
        # 处罚
        doc.add_paragraph(f"Penalty imposed: {employee_data.get('penalty_imposed', '')}")
        doc.add_paragraph(f"Reasons for penalty: {employee_data.get('penalty_reasons', '')}")
        
        # 上诉权
        doc.add_paragraph("Right of appeal: Employee has been advised of his / her right to appeal.")
        doc.add_paragraph(f"Employee: {employee_data.get('employee_name', '')}")
        doc.add_paragraph(
            "CCMA: The Employee was informed that should the case be referred to the Commission for Conciliation Mediation and Arbitration, such referral should take place within 30 days."
        )
        
        # 签名
        doc.add_paragraph(f"CHAIRPERSON (Signature): {'_'*20}")
        doc.add_paragraph(f"EMPLOYEE / WITNESS: {'_'*20}")
        
        # 确认
        doc.add_heading("ACKNOWLEDGEMENT THAT THE FINDINGS AND ACTIONS HAVE BEEN UNDERSTOOD", 2)
        doc.add_paragraph(f"EMPLOYEE / WITNESS: {'_'*20} DATE: {employee_data.get('acknowledgement_date', datetime.now().strftime('%d %B %Y'))}")
        doc.add_paragraph(f"CHAIRPERSON (Signature): {'_'*20} DATE: {employee_data.get('acknowledgement_date', datetime.now().strftime('%d %B %Y'))}")
        
        # 主席最终行动
        doc.add_heading("ACTION BY CHAIRPERSON", 2)
        final_actions = [
            "Inform complainant and Employee of the outcome of the case and the reasons for the penalty which has been imposed.",
            "The Employee must be advised of his / her right to appeal.",
            "The complainant and Employee must sign the disciplinary form (if the Employee refuses to sign, then a witness may sign in the presence of the Employee)."
        ]
        for i, action in enumerate(final_actions, 1):
            doc.add_paragraph(f"{i}. {action}")
        
        doc.add_paragraph(
            "NB: Signing of the document by the Employee does not imply acknowledgement of guilt. A copy must be handed to the Employee."
        )
        
        # 保存文档
        output_path = output_dir / f"{employee_data['employee_name']}_Disciplinary_Report.docx"
        doc.save(output_path)
        print(f"文档已生成: {output_path}")

    else:
        raise ValueError(f"Unsupported template type: {template_type}. Use 'B1', 'B2', 'E', 'F', 'G', 'H', 'C', or 'D'.")